#将docs\.pdf\.md\.txt文件统一化作中间格式:中转站
import re
from pathlib import Path

SUPPORTS_EXTS=[".docx",".txt",".pdf",".md",".xlsx"]

def parser_document(path:str)->list[dict]:
    docs=Path(path)
    text=docs.suffix.lower()  
    if text not in SUPPORTS_EXTS:
        raise ValueError("error")
    if text==".docx":
        return _parser_docx(path)
    if text==".pdf":
        return _parser_pdf(path)
    if text==".xlsx":
        return _parser_xlsx(path)
    return _parser_text(path)

# 以Document()打开对象,以段落格式保存

def _parser_docx(path:str)->list[dict]:
    from docx import Document
    doc=Document(path)
    segement:list[dict]=[]
    for x in doc.paragraphs:
        text=x.text.strip()
        if not text:
            continue
        style=(x.style.name or "").lower()
        m=re.search(r"heading\s*(\d+)",style)
        if m:
            segement.append({"type": "heading", "level": int(m.group(1)), "text": text})
        if not m:
            segement.append({"type":"text","level":0,"text":text})
    return segement

#以Pdf_reader() 打开对象,以行保存
def _parser_pdf(path:str)->list[dict]:
    from pypdf import PdfReader
    pdf=PdfReader(path)
    segement:list[dict]=[]
    for page in pdf.pages:
        text=(page.extract_text() or "").strip()
        for line in text.splitlines():
            lines=line.strip()
            if lines:
                segement.append({"type":"text","level":0,"text":lines})
    return segement

def _parser_text(path:str)->list[dict]:
    segements:list[dict]=[]
    with open(path,encoding="utf-8") as f:
        for line in f:
            line=line.strip()
            if not line:
                continue
            m=re.match(r"(#{1,4}\s+(.*)$)",line)
            if m:
                segements.append({"type":"heading","level":len(m.group(1)),"text":m.group(2)})
            else:
                segements.append({"type":"text","level":0,"text":line})
    return segements

def _parser_xlsx(path:str)->list[dict]:
    from openpyxl import load_workbook
    ws=load_workbook(path,data_only=True)
    segements:list[dict]=[]
    for w in ws.worksheets:
        rows=list(w.iter_rows(values_only=True)) #需要序列化 + 再拆分--》 list(迭代器)
        if not rows:
            continue
        head=["" if x is None else str(x).strip() for x in rows[0]] #读取序列0操作
        for i,row in enumerate(rows[1:],start=1):     #for 拆分,进行处理剩余行
            cells=["" if x is None else str(x).strip() for x in row]
            if not any(row):
                continue
            text=f"[表头]{"|".join(head)}\n[行{i}]:{"|".join(cells)}"
            segements.append({"type":"text","level":0,"text":text,
                              "meta":{"source_type":"table","sheet":w.title,"row_no":i,"header":head}})
    return segements

